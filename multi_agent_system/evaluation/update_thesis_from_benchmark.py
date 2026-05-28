#!/usr/bin/env python3
"""
Update thesis benchmark data (job74-83) menggantikan data lama (job54-67).

Usage:
  python3 multi_agent_system/evaluation/update_thesis_from_benchmark.py \
      --docx "Penelitian/Draft_70_Tugas Akhir_Latest_fixed.docx"

Jalankan dari root RAJDOLL setelah semua 10 scan benchmark selesai.
"""
import argparse
import json
import shutil
import statistics
import sys
from datetime import datetime
from pathlib import Path
from docx import Document

BASE_DIR = Path(__file__).parent
RUNS_DIR = BASE_DIR / "runs"
SUMMARY_FILE = BASE_DIR / "evaluation_juiceshop_summary.json"
RUN_LABELS = [f"juiceshop_benchmark_run{i}" for i in range(1, 11)]


def load_run_metrics(runs_dir: Path, run_label: str) -> dict:
    """Load metrics.json for one run, inject job_id from job_result.json."""
    metrics_file = runs_dir / run_label / "metrics.json"
    job_result_file = runs_dir / run_label / "job_result.json"
    if not metrics_file.exists():
        raise FileNotFoundError(f"Missing: {metrics_file}")
    m = json.loads(metrics_file.read_text())
    if job_result_file.exists():
        jr = json.loads(job_result_file.read_text())
        m["job_id"] = jr.get("job_id") or jr.get("id")
    return m


def compute_aggregate(metrics_list: list) -> dict:
    """Compute mean/std across all runs."""
    def mean_std(key):
        vals = [m[key] for m in metrics_list if m.get(key) is not None]
        if not vals:
            return None, None
        mu = round(statistics.mean(vals), 2)
        sd = round(statistics.pstdev(vals), 2) if len(vals) > 1 else 0.0
        return mu, sd

    p_mean, p_std = mean_std("precision")
    r_mean, r_std = mean_std("recall")
    f_mean, f_std = mean_std("f1")
    t_mean, t_std = mean_std("tcr")
    return {
        "n_runs": len(metrics_list),
        "precision_mean": p_mean, "precision_std": p_std,
        "recall_mean": r_mean,    "recall_std":    r_std,
        "f1_mean":     f_mean,    "f1_std":        f_std,
        "tcr_mean":    t_mean,    "tcr_std":       t_std,
    }


def update_summary_json(summary_file: Path, agg: dict, metrics_list: list, target: str = "juiceshop"):
    """Tulis ulang evaluation_juiceshop_summary.json dengan data baru."""
    per_run = []
    for m in metrics_list:
        entry = {k: v for k, v in m.items() if k != "job_id"}
        per_run.append(entry)
    data = {
        "n_runs": agg["n_runs"],
        "precision_mean": agg["precision_mean"],
        "precision_std":  agg["precision_std"],
        "recall_mean":    agg["recall_mean"],
        "recall_std":     agg["recall_std"],
        "f1_mean":        agg["f1_mean"],
        "f1_std":         agg["f1_std"],
        "tcr_mean":       agg["tcr_mean"],
        "tcr_std":        agg["tcr_std"],
        "scan_time_mean_hours": None,
        "scan_time_std_hours":  None,
        "per_run": per_run,
        "target": target,
    }
    summary_file.write_text(json.dumps(data, indent=2))


def _fmt(val: float) -> str:
    """Format float dengan koma desimal (Indonesian)."""
    return f"{val:.2f}".replace(".", ",")


def _fs(val: float) -> str:
    """Format std deviation with Indonesian comma, removing trailing zeros."""
    s = str(round(val, 2)).replace(".", ",")
    return s if s else "0"


def update_distribution_table(doc, metrics_list: list, agg: dict, table_index: int = 17):
    """Ganti 10 baris data + baris rata-rata di tabel distribusi run."""
    tbl = doc.tables[table_index]
    for i, m in enumerate(metrics_list, 1):
        job_id = m.get("job_id", f"run{i}")
        row = tbl.rows[i]
        row.cells[0].text = f"job{job_id}"
        row.cells[1].text = _fmt(m["precision"])
        row.cells[2].text = _fmt(m["recall"])
        row.cells[3].text = _fmt(m["f1"])
        row.cells[4].text = _fmt(m["tcr"])
        row.cells[5].text = ""

    avg_row = tbl.rows[11]
    avg_row.cells[0].text = "Rata-rata"
    avg_row.cells[1].text = f"{_fmt(agg['precision_mean'])} ±{_fs(agg['precision_std'])}"
    avg_row.cells[2].text = f"{_fmt(agg['recall_mean'])} ±{_fs(agg['recall_std'])}"
    avg_row.cells[3].text = f"{_fmt(agg['f1_mean'])} ±{_fs(agg['f1_std'])}"
    avg_row.cells[4].text = f"{_fmt(agg['tcr_mean'])} ±{_fs(agg['tcr_std'])}"
    avg_row.cells[5].text = "n=10 run Qwen3-4B"


def update_comparison_table(doc, baseline: dict, agg: dict, metrics_list: list, table_index: int = 22):
    """Update tabel perbandingan: job74 (Run Pertama) vs Rata-rata 10 Run."""
    tbl = doc.tables[table_index]
    job_id = baseline.get("job_id", 74)

    # Header row
    tbl.rows[0].cells[1].text = f"job{job_id} (Run Pertama)"
    tbl.rows[0].cells[2].text = "Rata-rata 10 Run (Final)"
    tbl.rows[0].cells[3].text = "Delta"

    total_baseline = baseline.get("total_findings_non_info", "?")
    total_mean = round(statistics.mean(
        m.get("total_findings_non_info", 0) for m in metrics_list
    ))
    tp_baseline = baseline.get("tp_gt_entries", "?")
    tp_mean = round(statistics.mean(m.get("tp_gt_entries", 0) for m in metrics_list))

    # Row 2: Total Temuan
    tbl.rows[2].cells[1].text = f"{total_baseline} temuan"
    tbl.rows[2].cells[2].text = f"~{total_mean} temuan (mean 10 run)"
    delta_temuan = total_mean - total_baseline if isinstance(total_baseline, int) else "?"
    tbl.rows[2].cells[3].text = f"+{delta_temuan}" if isinstance(delta_temuan, int) else "?"

    def delta_str(base, final):
        d = round(final - base, 2)
        sign = "+" if d >= 0 else ""
        return f"{sign}{_fmt(d)}%"

    # Row 3: Precision
    tbl.rows[3].cells[1].text = f"{_fmt(baseline['precision'])}%"
    tbl.rows[3].cells[2].text = f"{_fmt(agg['precision_mean'])}% (±{_fs(agg['precision_std'])}%)"
    tbl.rows[3].cells[3].text = delta_str(baseline['precision'], agg['precision_mean'])

    # Row 4: Recall
    tbl.rows[4].cells[1].text = f"{_fmt(baseline['recall'])}%"
    tbl.rows[4].cells[2].text = f"{_fmt(agg['recall_mean'])}% (±{_fs(agg['recall_std'])}%)"
    tbl.rows[4].cells[3].text = delta_str(baseline['recall'], agg['recall_mean'])

    # Row 5: F1-Score
    tbl.rows[5].cells[1].text = f"{_fmt(baseline['f1'])}%"
    tbl.rows[5].cells[2].text = f"{_fmt(agg['f1_mean'])}% (±{_fs(agg['f1_std'])}%)"
    tbl.rows[5].cells[3].text = delta_str(baseline['f1'], agg['f1_mean'])

    # Row 6: TCR
    tbl.rows[6].cells[1].text = f"{_fmt(baseline['tcr'])}%"
    tbl.rows[6].cells[2].text = f"{_fmt(agg['tcr_mean'])}% (±{_fs(agg['tcr_std'])}%)"
    tbl.rows[6].cells[3].text = delta_str(baseline['tcr'], agg['tcr_mean'])

    # Row 7: Tantangan Terdeteksi
    tbl.rows[7].cells[1].text = f"{tp_baseline}/57 entri GT"
    tbl.rows[7].cells[2].text = f"~{tp_mean}/57 entri GT (mean)"
    tbl.rows[7].cells[3].text = f"+{tp_mean - tp_baseline}" if isinstance(tp_baseline, int) else "?"

    # Row 8: Kegagalan Agen — biarkan (0, Stabil)


def update_text_paragraphs(doc, agg: dict, baseline_recall: float, baseline_job_id: int) -> list:
    """
    Ganti 9 pattern angka lama di paragraf teks.
    Urutan dari paling spesifik ke umum untuk mencegah partial match.
    Mengembalikan list string perubahan untuk logging.
    """
    def _fs2(v: float) -> str:
        return f"{v:.2f}".replace(".", ",")

    p_mean = _fmt(agg["precision_mean"])
    p_std  = _fs2(agg["precision_std"])
    r_mean = _fmt(agg["recall_mean"])
    r_std  = _fs2(agg["recall_std"])
    f_mean = _fmt(agg["f1_mean"])
    f_std  = _fs2(agg["f1_std"])
    t_mean = _fmt(agg["tcr_mean"])
    t_std  = _fs2(agg["tcr_std"])
    b_rec  = _fmt(baseline_recall)

    # Ordered: paling spesifik dulu
    patterns = [
        ("90,71% (±3,1%)",          f"{p_mean}% (±{p_std}%)"),
        ("92,10% (±4,38%)",         f"{r_mean}% (±{r_std}%)"),
        ("92,10% bukan 100%",       f"{r_mean}% bukan 100%"),
        ("91,40% (±3,73%)",         f"{f_mean}% (±{f_std}%)"),
        ("87,03% (±5,55%)",         f"{t_mean}% (±{t_std}%)"),
        ("68,42% (baseline run1)",  f"{b_rec}% (baseline job{baseline_job_id})"),
        ("91,23% (run9 final)",     f"{r_mean}% (rata-rata 10 run)"),
        ("91,40%",                  f"{f_mean}%"),
        ("87,03%",                  f"{t_mean}%"),
    ]

    changes = []
    for para in doc.paragraphs:
        for old, new in patterns:
            if old in para.text:
                replaced = False
                for run in para.runs:
                    if old in run.text:
                        run.text = run.text.replace(old, new)
                        changes.append(f"  '{old}' -> '{new}'")
                        replaced = True
                        break
                if not replaced:
                    # cross-run fallback
                    full = para.text
                    if old in full:
                        remainder = full.replace(old, new, 1)
                        for i, run in enumerate(para.runs):
                            run.text = remainder if i == 0 else ""
                        changes.append(f"  '{old}' -> '{new}' (cross-run)")
    return changes
