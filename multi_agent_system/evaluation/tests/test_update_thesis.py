import json, pytest
from pathlib import Path
from unittest.mock import patch, mock_open

# tambahkan path
import sys
sys.path.insert(0, str(Path(__file__).parent.parent))

from docx import Document
from update_thesis_from_benchmark import load_run_metrics, compute_aggregate, update_summary_json, update_distribution_table, update_comparison_table

SAMPLE_METRICS = {
    "precision": 90.54, "recall": 89.47, "f1": 89.99, "tcr": 85.19,
    "tp_gt_entries": 51, "fp_findings": 4, "fn_gt_entries": 6,
    "total_findings_non_info": 55, "total_gt_entries": 57,
    "run_label": "juiceshop_benchmark_run1"
}

def make_run_dir(tmp_path, run_name, metrics, job_id=74, job_result=None):
    d = tmp_path / "runs" / run_name
    d.mkdir(parents=True)
    (d / "metrics.json").write_text(json.dumps(metrics))
    if job_result is None:
        job_result = {"job_id": job_id, "status": "completed", "target": "http://juice-shop:3000/"}
    (d / "job_result.json").write_text(json.dumps(job_result))
    return d

def test_load_run_metrics_returns_metrics_and_job_id(tmp_path):
    runs_dir = tmp_path / "runs"
    make_run_dir(tmp_path, "juiceshop_benchmark_run1", SAMPLE_METRICS, job_id=74)
    result = load_run_metrics(runs_dir, "juiceshop_benchmark_run1")
    assert result["precision"] == 90.54
    assert result["job_id"] == 74

def test_load_run_metrics_missing_file_raises(tmp_path):
    runs_dir = tmp_path / "runs"
    with pytest.raises(FileNotFoundError):
        load_run_metrics(runs_dir, "juiceshop_benchmark_run1")

def test_compute_aggregate_mean_and_std():
    metrics_list = [
        {**SAMPLE_METRICS, "precision": 90.0, "recall": 88.0, "f1": 89.0, "tcr": 84.0},
        {**SAMPLE_METRICS, "precision": 92.0, "recall": 92.0, "f1": 92.0, "tcr": 88.0},
    ]
    agg = compute_aggregate(metrics_list)
    assert agg["precision_mean"] == pytest.approx(91.0)
    assert agg["recall_mean"] == pytest.approx(90.0)
    assert agg["precision_std"] == pytest.approx(1.0, abs=0.1)
    assert agg["n_runs"] == 2

def test_update_summary_json_writes_correct_data(tmp_path):
    summary_file = tmp_path / "evaluation_juiceshop_summary.json"
    metrics_list = [
        {**SAMPLE_METRICS, "precision": 90.0, "recall": 88.0, "f1": 89.0, "tcr": 84.0, "job_id": 74},
        {**SAMPLE_METRICS, "precision": 92.0, "recall": 92.0, "f1": 92.0, "tcr": 88.0, "job_id": 75},
    ]
    agg = compute_aggregate(metrics_list)
    update_summary_json(summary_file, agg, metrics_list, target="juiceshop")
    result = json.loads(summary_file.read_text())
    assert result["n_runs"] == 2
    assert result["precision_mean"] == pytest.approx(91.0)
    assert len(result["per_run"]) == 2
    assert result["target"] == "juiceshop"

def make_sample_docx(tmp_path):
    """Buat docx dengan tabel distribusi 12 baris (header + 10 data + rata-rata)."""
    doc = Document()
    tbl = doc.add_table(rows=12, cols=6)
    header = ['Run', 'P (%)', 'R (%)', 'F1 (%)', 'TCR (%)', 'Total']
    for j, h in enumerate(header):
        tbl.rows[0].cells[j].text = h
    old_runs = ['job54','job56','job57','job58','job59','job61','job64','job65','job66','job67']
    for i, r in enumerate(old_runs, 1):
        tbl.rows[i].cells[0].text = r
        tbl.rows[i].cells[1].text = '93,65'
        tbl.rows[i].cells[2].text = '96,49'
        tbl.rows[i].cells[3].text = '95,05'
        tbl.rows[i].cells[4].text = '92,59'
    tbl.rows[11].cells[0].text = 'Rata-rata'
    tbl.rows[11].cells[1].text = '90,71 ±3,1'
    path = tmp_path / "test.docx"
    doc.save(str(path))
    return path

def test_update_distribution_table_replaces_rows(tmp_path):
    metrics_list = [
        {**SAMPLE_METRICS, "precision": 91.0, "recall": 90.0, "f1": 90.5, "tcr": 85.0, "job_id": 74 + i}
        for i in range(10)
    ]
    agg = compute_aggregate(metrics_list)
    docx_path = make_sample_docx(tmp_path)
    doc = Document(str(docx_path))
    update_distribution_table(doc, metrics_list, agg, table_index=0)
    tbl = doc.tables[0]
    assert tbl.rows[1].cells[0].text == "job74"
    assert tbl.rows[10].cells[0].text == "job83"
    assert "91,00" in tbl.rows[11].cells[1].text  # rata-rata Precision

def make_comparison_docx(tmp_path):
    doc = Document()
    tbl = doc.add_table(rows=9, cols=4)
    rows_data = [
        ['Metrik', 'run1 (Baseline, Iterasi 3)', 'run9 (Final, Iterasi 4)', 'Delta'],
        ['Durasi Scan', '~37 menit', 'Di bawah 65 menit (mean 10 run Qwen)', '+75%'],
        ['Total Temuan', '62 temuan', '~71 temuan (mean 10 run Qwen)', '+14,5%'],
        ['Precision', '59,68%', '90,71% (±3,1%)', '+31,03%'],
        ['Recall', '68,42%', '92,10% (±4,38%)', '+23,68%'],
        ['F1-Score', '63,75%', '91,40% (±3,73%)', '+27,65%'],
        ['TCR', '62,96%', '87,03% (±5,55%)', '+24,07%'],
        ['Tantangan Terdeteksi', '39/57 entri GT', '~53/57 entri GT (mean)', '+13'],
        ['Kegagalan Agen', '0', '0', 'Stabil'],
    ]
    for i, row_data in enumerate(rows_data):
        for j, val in enumerate(row_data):
            tbl.rows[i].cells[j].text = val
    path = tmp_path / "comp.docx"
    doc.save(str(path))
    return path

def test_update_comparison_table_updates_header_and_metrics(tmp_path):
    baseline = {**SAMPLE_METRICS, "precision": 90.54, "recall": 89.47, "f1": 89.99,
                "tcr": 85.19, "tp_gt_entries": 51, "job_id": 74}
    agg = {
        "precision_mean": 91.0, "precision_std": 2.0,
        "recall_mean": 90.0, "recall_std": 3.0,
        "f1_mean": 90.5, "f1_std": 2.5,
        "tcr_mean": 86.0, "tcr_std": 4.0,
        "n_runs": 10,
    }
    metrics_list = [baseline] * 10
    docx_path = make_comparison_docx(tmp_path)
    doc = Document(str(docx_path))
    update_comparison_table(doc, baseline, agg, metrics_list, table_index=0)
    tbl = doc.tables[0]
    assert "job74" in tbl.rows[0].cells[1].text
    assert "Rata-rata" in tbl.rows[0].cells[2].text
    assert "90,54" in tbl.rows[3].cells[1].text   # precision baseline
    assert "91,00" in tbl.rows[3].cells[2].text   # precision final mean

def make_text_docx(tmp_path):
    doc = Document()
    doc.add_paragraph("Precision 90,71% (±3,1%), Recall 92,10% (±4,38%), dan F1-Score 91,40% semuanya melampaui.")
    doc.add_paragraph("TCR mencapai 87,03% (±5,55%) dengan waktu eksekusi.")
    doc.add_paragraph("Recall 92,10% bukan 100%.")
    doc.add_paragraph("91,40% standalone di kalimat lain.")
    doc.add_paragraph("87,03% standalone.")
    doc.add_paragraph("dari 68,42% (baseline run1) menjadi 91,23% (run9 final).")
    path = tmp_path / "text.docx"
    doc.save(str(path))
    return path

def test_update_text_paragraphs_replaces_all_patterns(tmp_path):
    from update_thesis_from_benchmark import update_text_paragraphs
    agg = {
        "precision_mean": 92.0, "precision_std": 2.5,
        "recall_mean": 93.0, "recall_std": 3.2,
        "f1_mean": 92.5, "f1_std": 2.8,
        "tcr_mean": 88.0, "tcr_std": 4.1,
    }
    baseline_recall = 89.47
    baseline_job_id = 74
    docx_path = make_text_docx(tmp_path)
    doc = Document(str(docx_path))
    changes = update_text_paragraphs(doc, agg, baseline_recall, baseline_job_id)
    texts = [p.text for p in doc.paragraphs]
    assert "92,00% (±2,50%)" in texts[0]   # precision with std
    assert "93,00% (±3,20%)" in texts[0]   # recall with std
    assert "92,50%" in texts[0]            # f1 standalone
    assert "88,00% (±4,10%)" in texts[1]   # tcr with std
    assert "93,00% bukan 100%" in texts[2]  # recall standalone pattern
    assert "88,00%" in texts[4]            # tcr standalone
    assert "89,47%" in texts[5]            # baseline recall
    assert "93,00% (rata-rata 10 run)" in texts[5]
    assert len(changes) >= 6
